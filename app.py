"""
Tesseract OCR Test Project
A simple Flask web application to test Tesseract OCR functionality.
"""

import os
import sys
import platform
from flask import Flask, render_template, request, jsonify, send_file
import pytesseract
from PIL import Image
import io
import tempfile
import shutil
from pdf2image import convert_from_path
from docx import Document

# === НАСТРОЙКА TESSERACT ДЛЯ WINDOWS ===
# Если система Windows, указываем путь к tesseract.exe явно
if platform.system() == 'Windows':
    # Стандартный путь установки Tesseract OCR в Windows
    tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    
    # Проверяем, существует ли файл по этому пути
    if os.path.exists(tesseract_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        print(f"[OK] Tesseract найден по пути: {tesseract_path}")
        TESSERACT_AVAILABLE = True
    else:
        # Если не найден в стандартном месте, пробуем альтернативное (для 32-bit или старых версий)
        alt_path = r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
        if os.path.exists(alt_path):
            pytesseract.pytesseract.tesseract_cmd = alt_path
            print(f"[OK] Tesseract найден по пути: {alt_path}")
            TESSERACT_AVAILABLE = True
        else:
            print(f"[ERROR] Tesseract не найден! Пожалуйста, установите Tesseract OCR.")
            print(f"Ожидаемый путь: {tesseract_path}")
            print(f"Или укажите правильный путь в коде app.py вручную.")
            TESSERACT_AVAILABLE = False
else:
    # Для Linux/Mac оставляем как есть (обычно добавлен в PATH)
    # Явно указываем путь к исполняемому файлу tesseract
    tesseract_cmd = '/usr/bin/tesseract'
    
    if os.path.isfile(tesseract_cmd):
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        TESSERACT_AVAILABLE = True
        print(f"[OK] Tesseract найден по пути: {tesseract_cmd}")
    else:
        # Пытаемся найти через PATH
        import shutil
        tesseract_cmd = shutil.which('tesseract')
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
            TESSERACT_AVAILABLE = True
            print(f"[OK] Tesseract найден через PATH: {tesseract_cmd}")
        else:
            TESSERACT_AVAILABLE = False
            print("[ERROR] Tesseract не найден в системе")

app = Flask(__name__)

# Проверка доступности Tesseract при старте
def check_tesseract_status():
    """Проверяет, доступен ли Tesseract и возвращает статус."""
    if not TESSERACT_AVAILABLE:
        return False, "Tesseract executable not found in system"
    try:
        version = pytesseract.get_tesseract_version()
        return True, f"Tesseract version: {version}"
    except Exception as e:
        return False, str(e)

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'}
ALLOWED_PDF_EXTENSIONS = {'pdf'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


def allowed_file(filename):
    """Check if the file has an allowed extension."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def allowed_pdf_file(filename):
    """Check if the file is a PDF."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_PDF_EXTENSIONS


def get_tesseract_version():
    """Get Tesseract version information."""
    if not TESSERACT_AVAILABLE:
        return "ERROR: tesseract is not installed or it's not in your PATH"
    try:
        version = pytesseract.get_tesseract_version()
        return str(version)
    except Exception as e:
        return f"Error: {str(e)}"


def get_available_languages():
    """Get list of available Tesseract languages."""
    if not TESSERACT_AVAILABLE:
        return ["ERROR: tesseract is not installed or it's not in your PATH"]
    try:
        langs = pytesseract.get_languages(config='')
        return langs
    except Exception as e:
        return [f"Error: {str(e)}"]


@app.route('/')
def index():
    """Render the main page."""
    tesseract_version = get_tesseract_version()
    available_languages = get_available_languages()
    # Default to both English and Russian
    default_lang = 'eng+rus'
    return render_template('index.html', 
                         tesseract_version=tesseract_version,
                         available_languages=available_languages,
                         default_lang=default_lang)


@app.route('/api/ocr', methods=['POST'])
def ocr_endpoint():
    """API endpoint for OCR processing."""
    if not TESSERACT_AVAILABLE:
        return jsonify({'error': 'Tesseract is not installed or not in your PATH'}), 503
        
    if 'image' not in request.files and 'image' not in request.form and 'pdf' not in request.files:
        return jsonify({'error': 'No image or PDF provided'}), 400
    
    try:
        # Handle PDF files
        if 'pdf' in request.files:
            file = request.files['pdf']
            if file.filename == '':
                return jsonify({'error': 'No selected file'}), 400
            
            if not allowed_pdf_file(file.filename):
                return jsonify({'error': 'File type not allowed. Only PDF files are accepted.'}), 400
            
            # Save PDF temporarily
            temp_dir = tempfile.mkdtemp()
            pdf_path = os.path.join(temp_dir, file.filename)
            file.save(pdf_path)
            
            # Get optional parameters
            lang = request.form.get('lang', 'eng+rus')
            psm = request.form.get('psm', None)
            
            # Build config string
            config = ''
            if psm:
                config += f' --psm {psm}'
            
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=300)
            
            # Process each page
            all_text = []
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image, lang=lang, config=config)
                all_text.append(f"--- Page {i+1} ---\n{page_text}")
            
            # Cleanup
            shutil.rmtree(temp_dir)
            
            text = '\n\n'.join(all_text)
            
            return jsonify({
                'success': True,
                'text': text,
                'language': lang,
                'pages_processed': len(images)
            })
        
        # Get image from form data (base64 or file upload)
        if 'image' in request.files:
            file = request.files['image']
            if file.filename == '':
                return jsonify({'error': 'No selected file'}), 400
            
            if not allowed_file(file.filename):
                return jsonify({'error': 'File type not allowed'}), 400
            
            image = Image.open(file.stream)
        else:
            # Handle base64 image data
            import base64
            image_data = request.form.get('image', '')
            if ',' in image_data:
                image_data = image_data.split(',')[1]
            
            image_bytes = base64.b64decode(image_data)
            image = Image.open(io.BytesIO(image_bytes))
        
        # Get optional parameters - default to eng+rus for bilingual support
        lang = request.form.get('lang', 'eng+rus')
        psm = request.form.get('psm', None)
        oem = request.form.get('oem', None)
        
        # Build config string
        config = ''
        if psm:
            config += f' --psm {psm}'
        if oem:
            config += f' --oem {oem}'
        
        # Perform OCR
        text = pytesseract.image_to_string(image, lang=lang, config=config)
        
        # Get detailed data if requested
        get_details = request.form.get('details', 'false').lower() == 'true'
        response = {
            'success': True,
            'text': text,
            'language': lang
        }
        
        if get_details:
            data = pytesseract.image_to_data(image, lang=lang, output_type=pytesseract.Output.DICT)
            response['details'] = {
                'confidence': sum([c for c in data['conf'] if c > 0]) / max(len([c for c in data['conf'] if c > 0]), 1),
                'words': len([t for t in data['text'] if t.strip()]),
                'lines': len(set(data['line_num']))
            }
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/test', methods=['GET'])
def test_endpoint():
    """Test endpoint to verify Tesseract is working with English and Russian."""
    try:
        from PIL import ImageDraw, ImageFont
        
        # Create a test image with both English and Russian text
        img = Image.new('RGB', (500, 150), color='white')
        draw = ImageDraw.Draw(img)
        
        # Try to use a font that supports Cyrillic
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
        except:
            try:
                font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", 24)
            except:
                font = ImageFont.load_default()
        
        # Add both English and Russian text
        english_text = "Hello World! OCR Test 123."
        russian_text = "Привет мир! Тест OCR 123."
        
        draw.text((10, 30), english_text, fill='black', font=font)
        draw.text((10, 70), russian_text, fill='black', font=font)
        
        # Perform OCR with both English and Russian languages
        text = pytesseract.image_to_string(img, lang='eng+rus')
        
        # Check if both languages are detected
        has_english = 'Hello' in text or 'World' in text or 'OCR' in text
        has_russian = 'Привет' in text or 'мир' in text or 'Тест' in text
        
        return jsonify({
            'success': True,
            'tesseract_version': get_tesseract_version(),
            'available_languages': get_available_languages(),
            'test_image_english': english_text,
            'test_image_russian': russian_text,
            'ocr_result': text.strip(),
            'english_detected': has_english,
            'russian_detected': has_russian,
            'both_languages_working': has_english and has_russian
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/export-docx', methods=['POST'])
def export_docx_endpoint():
    """API endpoint to export OCR result as Word document."""
    if not TESSERACT_AVAILABLE:
        return jsonify({'error': 'Tesseract is not installed or not in your PATH'}), 503
    
    try:
        data = request.get_json()
        text = data.get('text', '')
        language = data.get('language', 'eng+rus')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading('OCR Result', 0)
        
        # Add metadata
        doc.add_paragraph(f'Language: {language}')
        doc.add_paragraph(f'Date: {__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        # Add the recognized text
        # Split by lines and preserve structure
        lines = text.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()  # Empty line
        
        # Save to temporary file
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, 'ocr_result.docx')
        doc.save(docx_path)
        
        # Send file
        response = send_file(
            docx_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='ocr_result.docx'
        )
        
        # Cleanup after sending (note: this might not work perfectly with send_file)
        # For production, consider using a background task or different approach
        import atexit
        def cleanup():
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
        atexit.register(cleanup)
        
        return response
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    print("Starting Tesseract OCR Test Web Application...")
    print(f"Tesseract Version: {get_tesseract_version()}")
    print(f"Available Languages: {', '.join(get_available_languages())}")
    print("Open http://localhost:5000 in your browser")
    app.run(debug=True, host='0.0.0.0', port=5000)
